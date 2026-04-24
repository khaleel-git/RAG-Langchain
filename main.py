from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import subprocess
import textwrap
from pathlib import Path
from typing import Any, List, Optional

from dotenv import load_dotenv

try:
    from fpdf import FPDF
    from langchain_chroma import Chroma
    from langchain_core.documents import Document
    from langchain_core.prompts import ChatPromptTemplate
    from langchain_huggingface import HuggingFaceEmbeddings
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from pypdf import PdfReader
    from docx import Document as DocxDocument
except ImportError as exc:
    raise SystemExit(
        "Missing required packages. Install with:\n"
        "pip install -r requirements.txt\n"
        f"Import error: {exc}"
    )

try:
    from openai import OpenAI

    OPENAI_CLIENT_AVAILABLE = True
except ImportError:
    OPENAI_CLIENT_AVAILABLE = False
    OpenAI: Any = None

DEFAULT_NVIDIA_MODEL = "nvidia/llama-3.3-nemotron-super-49b-v1.5"
DEFAULT_NVIDIA_ATS_MODEL = "nvidia/nemotron-3-super-120b-a12b"
DEFAULT_NVIDIA_COVER_MODEL = "nvidia/nvidia-nemotron-nano-9b-v2"
DEFAULT_NVIDIA_BASE_URL = "https://integrate.api.nvidia.com/v1"
DEFAULT_JOB_DESCRIPTION_FILE = "job_description.txt"
DEFAULT_PDF_DIR = "Documents"
DEFAULT_PERSIST_DIR = "chroma_db"
DEFAULT_MATCH_FILE = "match.md"
DEFAULT_COVER_LETTER_FILE = "Khaleel_CoverLetter.pdf"
DEFAULT_COVER_LETTER_COPY_DIR = "output"
DEFAULT_RESUME_DOCX_INPUT = "Khaleel_Resume.docx"
DEFAULT_TAILORED_RESUME_DOCX = "khaleel_CV.docx"
DEFAULT_TAILORED_RESUME_DOC = "khaleel_CV.doc"

SKILL_ALIASES = {
    "Python": ["python", "pandas", "numpy", "pyspark"],
    "SQL": ["sql", "t-sql", "tsql", "pl/sql"],
    "MSSQL": ["mssql", "sql server", "ms sql"],
    "MySQL": ["mysql"],
    "PostgreSQL": ["postgresql", "postgres", "pgsql"],
    "MongoDB": ["mongodb", "mongo"],
    "Apache Airflow": ["apache airflow", "airflow", "dag"],
    "Docker": ["docker", "containerization"],
    "Kubernetes": ["kubernetes", "k8s"],
    "Terraform": ["terraform", "hcl"],
    "Ansible": ["ansible"],
    "Linux": ["linux", "ubuntu", "debian", "centos", "bash", "zsh"],
    "Git": ["git", "github", "gitlab", "bitbucket"],
    "CI/CD": ["ci/cd", "cicd", "github actions", "pipeline"],
    "Power BI": ["power bi", "powerbi"],
    "AWS": ["aws", "amazon web services"],
    "GCP": ["gcp", "google cloud"],
    "Azure": ["azure", "microsoft azure"],
    "VMware": ["vmware", "vcenter"],
    "Prometheus": ["prometheus"],
    "Grafana": ["grafana"],
    "Flask": ["flask"],
    "Django": ["django"],
    "React": ["react", "reactjs"],
    "Prompt Engineering": ["prompt engineering", "prompt tuning", "llm prompts"],
    "MLOps": ["mlops", "machine learning operations"],
    "Microsoft Office": ["microsoft office", "excel", "powerpoint", "word"],
    "German": ["german", "deutsch"],
    "English": ["english"],
}


def _skill_alias_regex(alias: str) -> str:
    escaped = re.escape(alias)
    if alias.isalnum():
        return rf"(?<!\w){escaped}(?!\w)"
    return escaped


def load_environment(env_file: str) -> None:
    env_path = Path(env_file)
    if not env_path.is_absolute():
        env_path = Path.cwd() / env_path

    if env_path.exists():
        load_dotenv(dotenv_path=env_path, override=True)
    else:
        load_dotenv(override=False)


def normalize_profile_text_for_indexing(text: str) -> str:
    return " ".join(text.lower().split())


def clean_job_description_text(text: str) -> str:
    text = text.replace("\u00ad", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\u0000-\u001f\u007f]", " ", text)
    text = re.sub(r"[•·●▪◦■]", " ", text)
    text = re.sub(r"-\s*\n\s*", "", text)
    text = re.sub(r"\n+", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def find_pdfs(pdf_dir: Path) -> List[Path]:
    return sorted(
        [
            path
            for path in pdf_dir.rglob("*")
            if path.is_file() and path.suffix.lower() == ".pdf"
        ]
    )


def load_pdf_documents(pdf_dir: Path) -> List[Document]:
    pdf_files = find_pdfs(pdf_dir)
    if not pdf_files:
        raise ValueError(f"No PDF files found under: {pdf_dir}")

    docs: List[Document] = []
    for pdf_path in pdf_files:
        reader = PdfReader(str(pdf_path))
        for page_idx, page in enumerate(reader.pages, start=1):
            text = normalize_profile_text_for_indexing((page.extract_text() or "").strip())
            if not text:
                continue
            docs.append(
                Document(
                    page_content=text,
                    metadata={"source": str(pdf_path), "page": page_idx},
                )
            )

    if not docs:
        raise ValueError(f"PDF files found, but no extractable text under: {pdf_dir}")

    return docs


def load_job_description(job_description_path: str) -> List[Document]:
    path = Path(job_description_path)
    if not path.is_absolute():
        path = Path.cwd() / path

    if not path.exists():
        raise FileNotFoundError(f"Job description file not found: {path}")

    if path.suffix.lower() == ".pdf":
        reader = PdfReader(str(path))
        docs: List[Document] = []
        for page_idx, page in enumerate(reader.pages, start=1):
            text = clean_job_description_text((page.extract_text() or "").strip())
            if not text:
                continue
            docs.append(
                Document(
                    page_content=text,
                    metadata={"source": str(path), "page": page_idx, "type": "job_description"},
                )
            )
        if not docs:
            raise ValueError(f"Job description PDF has no extractable text: {path}")
        return docs

    text = clean_job_description_text(path.read_text(encoding="utf-8").strip())
    if not text:
        raise ValueError(f"Job description file is empty: {path}")
    return [Document(page_content=text, metadata={"source": str(path), "type": "job_description"})]


def build_clean_vectorstore(
    docs: List[Document],
    persist_directory: Path,
    chunk_size: int,
    chunk_overlap: int,
    rebuild: bool,
) -> Chroma:
    if chunk_size <= 0:
        raise ValueError("chunk_size must be > 0")
    if chunk_overlap < 0:
        raise ValueError("chunk_overlap must be >= 0")
    if chunk_overlap >= chunk_size:
        raise ValueError("chunk_overlap must be smaller than chunk_size")

    persist_directory = persist_directory.resolve()
    workspace_root = Path.cwd().resolve()
    if workspace_root not in [persist_directory, *persist_directory.parents]:
        raise ValueError(f"Refusing to use persist directory outside workspace: {persist_directory}")
    if persist_directory == workspace_root:
        raise ValueError("Refusing to use workspace root as persist directory")

    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

    if persist_directory.exists() and not rebuild:
        try:
            return Chroma(
                embedding_function=embeddings,
                persist_directory=str(persist_directory),
            )
        except Exception:
            # If existing index is corrupted/incompatible, rebuild it.
            shutil.rmtree(persist_directory)
    elif persist_directory.exists() and rebuild:
        shutil.rmtree(persist_directory)

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    splits = text_splitter.split_documents(docs)

    return Chroma.from_documents(
        documents=splits,
        embedding=embeddings,
        persist_directory=str(persist_directory),
    )


def find_case_insensitive_skill_matches(profile_docs: List[Document], job_description_text: str) -> List[str]:
    profile_text = "\n".join(doc.page_content for doc in profile_docs).lower()
    jd_text = job_description_text.lower()

    matches: List[str] = []
    for canonical_skill, aliases in SKILL_ALIASES.items():
        present_in_jd = any(re.search(_skill_alias_regex(alias), jd_text) for alias in aliases)
        present_in_profile = any(re.search(_skill_alias_regex(alias), profile_text) for alias in aliases)
        if present_in_jd and present_in_profile:
            matches.append(canonical_skill)
    return matches


def build_job_match_prompt() -> ChatPromptTemplate:
    return ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "You are an ATS-style resume matcher. Compare candidate profile context against the job description. "
                "Do not invent experience. Return concise structured output.",
            ),
            (
                "human",
                "Job description:\n{input}\n\nCandidate profile context:\n{context}\n\n"
                "Return these sections:\n"
                "1. Overall score as an integer from 0 to 100.\n"
                "2. Strengths and gaps.\n"
                "3. Top matching keywords or skills.\n"
                "4. Recommendation for improving the match.",
            ),
        ]
    )


def _to_openai_messages(messages: List[Any]) -> List[dict[str, str]]:
    converted: List[dict[str, str]] = []
    for msg in messages:
        role = getattr(msg, "type", "user")
        content = str(getattr(msg, "content", ""))
        if role == "human":
            openai_role = "user"
        elif role == "ai":
            openai_role = "assistant"
        elif role == "system":
            openai_role = "system"
        else:
            openai_role = "user"
        converted.append({"role": openai_role, "content": content})
    return converted


def _resolve_provider(provider: str) -> str:
    if provider in {"auto", "nvidia"}:
        return provider
    return "nvidia"


def _invoke_chat(
    messages: List[Any],
    model: Optional[str],
    provider: str,
    temperature: float,
    task: str,
) -> str:
    resolved_provider = _resolve_provider(provider)

    if resolved_provider == "nvidia":
        if not OPENAI_CLIENT_AVAILABLE:
            raise RuntimeError("openai package is missing. Install with pip install openai")

        if task == "ats":
            nvidia_api_key = os.getenv("NVIDIA_ATS_API_KEY") or os.getenv("NVIDIA_API_KEY")
        elif task == "cover":
            nvidia_api_key = os.getenv("NVIDIA_COVER_API_KEY") or os.getenv("NVIDIA_API_KEY")
        else:
            nvidia_api_key = os.getenv("NVIDIA_API_KEY")

        if not nvidia_api_key:
            if task == "ats":
                raise RuntimeError(
                    "Missing NVIDIA API key for ATS. Set NVIDIA_ATS_API_KEY or NVIDIA_API_KEY in environment."
                )
            if task == "cover":
                raise RuntimeError(
                    "Missing NVIDIA API key for cover letter. Set NVIDIA_COVER_API_KEY or NVIDIA_API_KEY in environment."
                )
            raise RuntimeError("NVIDIA_API_KEY is not set in environment.")

        assert OpenAI is not None
        client = OpenAI(
            base_url=os.getenv("NVIDIA_BASE_URL", DEFAULT_NVIDIA_BASE_URL),
            api_key=nvidia_api_key,
        )
        if task == "ats":
            default_task_model = DEFAULT_NVIDIA_ATS_MODEL
            env_task_model = os.getenv("NVIDIA_ATS_MODEL")
        elif task == "cover":
            default_task_model = DEFAULT_NVIDIA_COVER_MODEL
            env_task_model = os.getenv("NVIDIA_COVER_MODEL")
        else:
            default_task_model = DEFAULT_NVIDIA_MODEL
            env_task_model = None

        selected_model = model or env_task_model or os.getenv("NVIDIA_MODEL") or default_task_model

        request_kwargs: dict[str, Any] = {
            "model": selected_model,
            "messages": _to_openai_messages(messages),
            "temperature": temperature,
            "top_p": 0.95,
            "max_tokens": 4096,
            "stream": False,
        }

        if task == "ats":
            request_kwargs["max_tokens"] = 16384
            request_kwargs["extra_body"] = {
                "reasoning_budget": 16384,
                "chat_template_kwargs": {"enable_thinking": True},
            }
        elif task == "cover":
            request_kwargs["max_tokens"] = 2048

        try:
            response = client.chat.completions.create(**request_kwargs)
        except Exception:
            # Retry once without provider-specific extra body for broader compatibility.
            request_kwargs.pop("extra_body", None)
            response = client.chat.completions.create(**request_kwargs)

        content = response.choices[0].message.content
        return content.strip() if content else ""

    raise RuntimeError(f"Unsupported provider: {resolved_provider}")


def invoke_llm(
    prompt: ChatPromptTemplate,
    query: str,
    context_docs: List[Document],
    model: Optional[str],
    provider: str,
) -> str:
    context = "\n\n".join(doc.page_content for doc in context_docs)
    messages = prompt.format_messages(input=query, context=context)
    return _invoke_chat(messages, model=model, provider=provider, temperature=0.0, task="ats")


def format_match_markdown(
    job_description_text: str,
    matched_skills: List[str],
    evidence_docs: List[Document],
    answer: str,
    answer_is_fallback: bool,
) -> str:
    skills_text = ", ".join(matched_skills) if matched_skills else "None"
    lines: List[str] = [
        "# Job Match Report",
        "",
        "## Question",
        "",
        job_description_text,
        "",
        "## Case-Insensitive Skill Overlaps",
        "",
        skills_text,
        "",
        "## Retrieved Resume Evidence",
        "",
    ]

    for idx, doc in enumerate(evidence_docs, start=1):
        source = doc.metadata.get("source", "unknown")
        content = " ".join(doc.page_content.split())
        lines.extend([f"### Chunk {idx}", "", f"- Source: {source}", f"- Text: {content}", ""])

    lines.extend(
        [
            "## Result",
            "",
            "(Fallback mode: retrieved context only)" if answer_is_fallback else "(LLM generated)",
            "",
            answer,
            "",
        ]
    )
    return "\n".join(lines)


def print_formatted_match_report(
    job_description_text: str,
    matched_skills: List[str],
    evidence_docs: List[Document],
    answer: str,
    answer_is_fallback: bool,
) -> None:
    print("\n" + "=" * 80)
    print("ATS SCORE REPORT")
    print("=" * 80)
    print("\nQ:")
    print(textwrap.fill(job_description_text, width=100))

    skills_text = ", ".join(matched_skills) if matched_skills else "None"
    print(f"\nCase-insensitive skill overlaps: {skills_text}")

    print("\nResume evidence:")
    for idx, doc in enumerate(evidence_docs, start=1):
        source = doc.metadata.get("source", "unknown")
        content = " ".join(doc.page_content.split())
        print(f"\n[{idx}] {source}")
        print(textwrap.fill(content, width=100))

    print("\nResult:")
    if answer_is_fallback:
        print("(Fallback mode: retrieved context only)")
    print(answer)
    print("=" * 80)


def extract_section(markdown: str, heading: str, next_heading: Optional[str] = None) -> str:
    start_pattern = re.escape(heading)
    if next_heading:
        end_pattern = re.escape(next_heading)
        pattern = rf"{start_pattern}\s*(.*?){end_pattern}"
    else:
        pattern = rf"{start_pattern}\s*(.*)$"

    match = re.search(pattern, markdown, flags=re.IGNORECASE | re.DOTALL)
    if not match:
        return ""
    return match.group(1).strip()


def extract_overall_score(markdown: str) -> int:
    # Flexible regex to find "Overall score: 58" or "Overall score: 58 / 100"
    match = re.search(r"overall\s*score[^\d]*(\d{1,3})", markdown, flags=re.IGNORECASE)
    if not match:
        # Fallback for "Score: 58"
        match = re.search(r"\bscore[^\d]*(\d{1,3})", markdown, flags=re.IGNORECASE)

    if not match:
        return 0  # Return 0 if no score is found, instead of raising an error.

    return int(match.group(1))


def parse_skill_overlaps(markdown: str) -> List[str]:
    section = extract_section(markdown, "## Case-Insensitive Skill Overlaps", "## Retrieved Resume Evidence")
    if not section:
        return []
    return [item.strip() for item in section.split(",") if item.strip()]


def is_job_german_eligible(question_text: str) -> bool:
    text = question_text.lower()

    language_context_patterns = [
        r"\b(german|deutsch)\s+(language|skills?|knowledge|proficiency|level|spoken|written|fluency)\b",
        r"\b(language|skills?|knowledge|proficiency|level|spoken|written|fluency)\s+(in\s+)?(german|deutsch)\b",
        r"\b(german|deutsch)\s+and\s+english\s+skills\b",
        r"\b(excellent|fluent|very\s+good|proficient|basic|elementary|limited|beginner)\s+(german|deutsch)\b",
    ]
    has_language_mention = any(re.search(pattern, text) for pattern in language_context_patterns)
    if not has_language_mention:
        return True

    basic_patterns = [
        r"basic\s+german",
        r"elementary\s+german",
        r"limited\s+german",
        r"beginner\s+german",
        r"basic\s+deutsch",
        r"elementary\s+deutsch",
        r"\ba1\b",
        r"\ba2\b",
        r"\bb1\b",
    ]

    advanced_patterns = [
        r"excellent\s+german",
        r"fluent\s+german",
        r"very\s+good\s+german",
        r"proficient\s+german",
        r"native\s+german",
        r"excellent\s+deutsch",
        r"flie(?:ss|ß)end",
        r"\bc1\b",
        r"\bc2\b",
    ]

    if any(re.search(pattern, text) for pattern in advanced_patterns):
        return False
    return any(re.search(pattern, text) for pattern in basic_patterns)


def generate_cover_letter_text(
    job_text: str,
    matched_skills: List[str],
    recommendations: str,
    chunks: List[Document],
    cover_model: Optional[str],
    provider: str,
) -> str:
    context = "\n\n".join(doc.page_content for doc in chunks)
    skills = ", ".join(matched_skills[:4]) if matched_skills else "data, reporting, and automation"

    try:
        prompt = ChatPromptTemplate.from_messages(
            [
                (
                    "system",
                    "You are a professional career assistant. Write like a real person. "
                    "Use simple English and short sentences. Keep it clear and human-readable.",
                ),
                (
                    "human",
                    "Write a one-page cover letter for Khaleel Ahmad based on this context.\n\n"
                    "Job description:\n{job}\n\n"
                    "Matched skills:\n{skills}\n\n"
                    "Recommendations:\n{recommendations}\n\n"
                    "Resume evidence:\n{context}\n\n"
                    "Rules:\n"
                    "1. Tone must be natural and human.\n"
                    "2. Use simple English.\n"
                    "3. Mention willingness to relocate.\n"
                    "4. Mention ongoing German B1 learning.\n"
                    "5. Do not invent achievements.\n"
                    "6. Output only the final cover letter with greeting and closing.",
                ),
            ]
        )

        messages = prompt.format_messages(
            job=job_text,
            skills=skills,
            recommendations=recommendations,
            context=context,
        )
        response_text = _invoke_chat(
            messages,
            model=cover_model,
            provider=provider,
            temperature=0.2,
            task="cover",
        )
        if response_text:
            return response_text
    except Exception:
        pass

    first_para = (
        "Dear Hiring Manager,\n\n"
        "I am applying for this role because it matches my background and goals. "
        f"My strongest matching skills are {skills}."
    )
    second_para = (
        "\n\nI have worked on practical data and automation tasks, and I enjoy building useful solutions for teams."
        " I have hands-on experience from my recent work and study projects. "
        "I am currently studying Master of Science in Artificial Intelligence in Germany."
    )
    third_para = (
        "\n\nI am willing to relocate for this position. I am also doing German B1 to improve my communication at work. "
        "Thank you for your time and consideration.\n\n"
        "Sincerely,\n"
        "Khaleel Ahmad"
    )
    return first_para + second_para + third_para


def generate_resume_targeting_content(
    job_text: str,
    recommendations: str,
    chunks: List[Document],
    resume_model: Optional[str],
    provider: str,
) -> dict[str, str]:
    content = {
        "summary": (
            "M.Sc. Artificial Intelligence student in Germany seeking a Werkstudent role in AI tooling and "
            "consulting support, with hands-on data engineering experience in Python, SQL, and automation."
        ),
        "skills_suffix": "Prompt Engineering (LLM testing), AI Tool Validation, Stakeholder Communication",
        "office_suffix": "Microsoft Office (Excel, PowerPoint, Word)",
        "language_suffix": "English: Professional, German: B1 (in progress)",
    }

    context = "\n\n".join(doc.page_content for doc in chunks[:6])
    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "You are an expert resume tailoring assistant. Keep outputs factual, concise, and professional.",
            ),
            (
                "human",
                "Tailor resume snippets for this job using only provided evidence.\n\n"
                "Job description:\n{job}\n\n"
                "Recommendations:\n{recommendations}\n\n"
                "Resume evidence:\n{context}\n\n"
                "Return JSON only with exactly these keys: summary, skills_suffix, office_suffix, language_suffix.\n"
                "Constraints:\n"
                "1. Do not invent achievements, companies, or certifications.\n"
                "2. Keep each field under 20 words.\n"
                "3. Avoid location/company names in skills.\n"
                "4. Keep language level wording conservative if uncertain.",
            ),
        ]
    )

    try:
        messages = prompt.format_messages(job=job_text, recommendations=recommendations, context=context)
        response_text = _invoke_chat(
            messages,
            model=resume_model,
            provider=provider,
            temperature=0.1,
            task="cover",
        )

        json_match = re.search(r"\{.*\}", response_text, flags=re.DOTALL)
        if json_match:
            parsed = json.loads(json_match.group(0))
            if isinstance(parsed, dict):
                for key in content:
                    value = parsed.get(key)
                    if isinstance(value, str) and value.strip():
                        sanitized = " ".join(value.strip().split())
                        content[key] = sanitized[:180]
    except Exception:
        pass

    combined_text = f"{job_text}\n{recommendations}".lower()
    if "prompt" not in combined_text:
        content["skills_suffix"] = "AI Tool Validation, Stakeholder Communication"
    if "office" not in combined_text and "microsoft" not in combined_text:
        content["office_suffix"] = "Microsoft Office"
    if "german" not in combined_text and "deutsch" not in combined_text:
        content["language_suffix"] = "English: Professional"
    return content


def _append_suffix_if_missing(line: str, suffix: str) -> str:
    normalized_line = line.lower()
    normalized_suffix = suffix.lower()
    if normalized_suffix in normalized_line:
        return line
    return f"{line}, {suffix}" if line.strip() else suffix


def tailor_resume_docx(
    input_docx: Path,
    output_docx: Path,
    targeting_content: dict[str, str],
) -> bool:
    if not input_docx.exists():
        raise FileNotFoundError(f"Resume template not found: {input_docx}")

    shutil.copy2(input_docx, output_docx)
    doc = DocxDocument(str(output_docx))

    summary = targeting_content.get("summary", "").strip()
    skills_suffix = targeting_content.get("skills_suffix", "").strip()
    office_suffix = targeting_content.get("office_suffix", "").strip()
    language_suffix = targeting_content.get("language_suffix", "").strip()

    changed = False
    inserted_summary = False

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        lower_text = text.lower()

        if text and "|" in text and "@" in text and not inserted_summary and summary:
            insert_idx = idx + 1
            new_paragraph = doc.paragraphs[insert_idx].insert_paragraph_before(summary)
            if paragraph.style is not None:
                new_paragraph.style = paragraph.style
            inserted_summary = True
            changed = True
            continue

        if lower_text.startswith("languages:") and language_suffix:
            updated = _append_suffix_if_missing(text, language_suffix)
            if updated != text:
                paragraph.text = updated
                changed = True
            continue

        if lower_text.startswith("data engineering:") and skills_suffix:
            updated = _append_suffix_if_missing(text, skills_suffix)
            if updated != text:
                paragraph.text = updated
                changed = True
            continue

        if lower_text.startswith("frameworks") and office_suffix:
            updated = _append_suffix_if_missing(text, office_suffix)
            if updated != text:
                paragraph.text = updated
                changed = True

    if changed:
        doc.save(str(output_docx))
    return changed


def convert_docx_to_doc(input_docx: Path, output_doc: Path) -> None:
    command = [
        "textutil",
        "-convert",
        "doc",
        "-output",
        str(output_doc),
        str(input_docx),
    ]
    subprocess.run(command, check=True, capture_output=True, text=True)


def save_cover_letter_pdf(letter_text: str, output_path: Path) -> None:
    def sanitize_pdf_text(text: str) -> str:
        text = text.replace("\u00ad", "")
        text = text.replace("\u200b", "")
        text = text.replace("\u2013", "-").replace("\u2014", "-")
        text = text.replace("\u2018", "'").replace("\u2019", "'")
        text = text.replace("\u201c", '"').replace("\u201d", '"')
        text = re.sub(r"(\S{60})(?=\S)", r"\1 ", text)
        return text.encode("latin-1", "replace").decode("latin-1")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)

    for paragraph in letter_text.split("\n"):
        safe_paragraph = sanitize_pdf_text(paragraph)
        if safe_paragraph:
            pdf.multi_cell(0, 8, safe_paragraph)
        else:
            pdf.ln(4)
        pdf.ln(1)

    pdf.output(str(output_path))


def copy_cover_letter_to_target(output_path: Path, target_dir: Path) -> Path:
    if not output_path.exists():
        raise FileNotFoundError(f"Generated cover letter not found: {output_path}")

    target_dir.mkdir(parents=True, exist_ok=True)
    target_file = target_dir / output_path.name
    shutil.copy2(output_path, target_file)
    return target_file


def remove_stale_cover_letter_files(output_path: Path, target_dir: Path) -> None:
    stale_files = [output_path, target_dir / output_path.name]
    for stale_file in stale_files:
        try:
            if stale_file.exists():
                stale_file.unlink()
        except Exception:
            # Best-effort cleanup only; pipeline should continue even if deletion fails.
            pass


def run_ats_and_cover_pipeline(args: argparse.Namespace) -> None:
    load_environment(args.env_file)

    resume_docx_input_env = os.getenv("RESUME_DOCX_INPUT")
    resume_docx_output_env = os.getenv("RESUME_DOCX_OUTPUT")
    resume_doc_output_env = os.getenv("RESUME_DOC_OUTPUT")

    pdf_dir = Path(args.pdf_dir)
    if not pdf_dir.is_absolute():
        pdf_dir = Path.cwd() / pdf_dir

    persist_dir = Path(args.persist_dir)
    if not persist_dir.is_absolute():
        persist_dir = Path.cwd() / persist_dir

    match_file = Path(args.match_file)
    if not match_file.is_absolute():
        match_file = Path.cwd() / match_file

    cover_letter_output = Path(args.cover_letter_output)
    if not cover_letter_output.is_absolute():
        cover_letter_output = Path.cwd() / cover_letter_output

    resume_docx_input = Path(resume_docx_input_env or args.resume_docx_input)
    if not resume_docx_input.is_absolute():
        resume_docx_input = Path.cwd() / resume_docx_input

    tailored_resume_docx_output = Path(resume_docx_output_env or args.resume_docx_output)
    if not tailored_resume_docx_output.is_absolute():
        tailored_resume_docx_output = Path.cwd() / tailored_resume_docx_output

    tailored_resume_doc_output = Path(resume_doc_output_env or args.resume_doc_output)
    if not tailored_resume_doc_output.is_absolute():
        tailored_resume_doc_output = Path.cwd() / tailored_resume_doc_output

    profile_docs = load_pdf_documents(pdf_dir)
    print(f"Loaded {len(profile_docs)} profile pages from PDF files at: {pdf_dir}")

    job_docs = load_job_description(args.job_description)
    job_text = clean_job_description_text("\n\n".join(doc.page_content for doc in job_docs).strip())
    if not job_text:
        raise ValueError("Job description text is empty.")
    print(f"Loaded job description from: {args.job_description}")

    vectorstore = build_clean_vectorstore(
        docs=profile_docs,
        persist_directory=persist_dir,
        chunk_size=args.chunk_size,
        chunk_overlap=args.chunk_overlap,
        rebuild=args.rebuild,
    )

    matched_skills = find_case_insensitive_skill_matches(profile_docs, job_text)
    overlap_text = ", ".join(matched_skills) if matched_skills else "No direct skill overlaps found in configured skill list"

    score_query = (
        f"{job_text}\n\n"
        f"Case-insensitive skill scan across all profile documents: {overlap_text}.\n"
        "Treat these as explicit matches from the resume corpus."
    )

    score_retriever = vectorstore.as_retriever(search_kwargs={"k": max(args.k, 8)})
    score_evidence = score_retriever.invoke(score_query)

    prompt = build_job_match_prompt()
    ats_model = args.ats_model or args.model
    cover_model = args.cover_model or args.model
    try:
        score_answer = invoke_llm(prompt, score_query, score_evidence, ats_model, args.provider)
        score_is_fallback = False
    except Exception as exc:
        score_answer = f"Failed to generate ATS assessment with selected provider: {exc}"
        score_is_fallback = True

    print_formatted_match_report(job_text, matched_skills, score_evidence, score_answer, score_is_fallback)

    try:
        score = extract_overall_score(score_answer)
    except ValueError:
        score = 0

    match_markdown = format_match_markdown(
        job_description_text=job_text,
        matched_skills=matched_skills,
        evidence_docs=score_evidence,
        answer=score_answer,
        answer_is_fallback=score_is_fallback,
    )
    match_file.write_text(match_markdown, encoding="utf-8")
    print(f"\nSaved formatted match report to: {match_file}")

    copy_dir_from_env = os.getenv("COVER_LETTER_COPY_DIR")
    cover_letter_copy_dir = Path(copy_dir_from_env or args.cover_letter_copy_dir)
    if not cover_letter_copy_dir.is_absolute():
        cover_letter_copy_dir = Path.cwd() / cover_letter_copy_dir

    if score <= 50:
        remove_stale_cover_letter_files(cover_letter_output, cover_letter_copy_dir)
        print(f"\nCover letter not generated. Overall score is {score}, must be > 50.")
        print("Removed stale cover-letter PDF (if present) to avoid mismatch with current report.")
        return

    if not is_job_german_eligible(job_text):
        remove_stale_cover_letter_files(cover_letter_output, cover_letter_copy_dir)
        print("\nCover letter not generated. Job description has German-language requirement above basic level.")
        print("Removed stale cover-letter PDF (if present) to avoid mismatch with current report.")
        return

    recommendations = extract_section(match_markdown, "## Result", None)
    cover_retriever = vectorstore.as_retriever(search_kwargs={"k": max(args.k, 8)})
    cover_query = (
        f"Job description:\n{job_text}\n\n"
        f"Recommendations:\n{recommendations}\n\n"
        "Find best resume evidence for writing a targeted cover letter."
    )
    cover_chunks = cover_retriever.invoke(cover_query)

    letter_text = generate_cover_letter_text(
        job_text=job_text,
        matched_skills=parse_skill_overlaps(match_markdown),
        recommendations=recommendations,
        chunks=cover_chunks,
        cover_model=cover_model,
        provider=args.provider,
    )
    save_cover_letter_pdf(letter_text, cover_letter_output)
    print(f"\nCover letter generated: {cover_letter_output}")

    copied_cover_letter = copy_cover_letter_to_target(
        cover_letter_output,
        cover_letter_copy_dir,
    )
    print(f"Cover letter copied to: {copied_cover_letter}")

    resume_targeting_content = generate_resume_targeting_content(
        job_text=job_text,
        recommendations=recommendations,
        chunks=cover_chunks,
        resume_model=cover_model,
        provider=args.provider,
    )

    resume_updated = tailor_resume_docx(
        input_docx=resume_docx_input,
        output_docx=tailored_resume_docx_output,
        targeting_content=resume_targeting_content,
    )
    if resume_updated:
        print(f"Tailored DOCX resume generated: {tailored_resume_docx_output}")
    else:
        print("Tailored resume generation ran, but no section was changed in DOCX template.")

    try:
        convert_docx_to_doc(tailored_resume_docx_output, tailored_resume_doc_output)
        print(f"Tailored DOC resume generated: {tailored_resume_doc_output}")
    except Exception as exc:
        print(
            "Warning: Failed to convert tailored DOCX to DOC with textutil. "
            f"DOCX is available at {tailored_resume_docx_output}. Details: {exc}"
        )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="ATS score calculator + cover letter generator")
    parser.add_argument(
        "--job-description",
        dest="job_description",
        default=DEFAULT_JOB_DESCRIPTION_FILE,
        help="Path to job description file (.txt/.pdf). Default: job_description.txt",
    )
    parser.add_argument(
        "--pdf-dir",
        dest="pdf_dir",
        default=DEFAULT_PDF_DIR,
        help="Folder containing profile PDFs. Default: Documents",
    )
    parser.add_argument(
        "--persist-dir",
        dest="persist_dir",
        default=DEFAULT_PERSIST_DIR,
        help="Chroma persistence directory. Default: chroma_db",
    )
    parser.add_argument(
        "--match-file",
        default=DEFAULT_MATCH_FILE,
        help="Output markdown report path. Default: match.md",
    )
    parser.add_argument(
        "--cover-letter-output",
        default=DEFAULT_COVER_LETTER_FILE,
        help="Output PDF path for cover letter. Default: Khaleel_CoverLetter.pdf",
    )
    parser.add_argument(
        "--cover-letter-copy-dir",
        default=DEFAULT_COVER_LETTER_COPY_DIR,
        help="Directory to copy generated cover letter. Default: output (or COVER_LETTER_COPY_DIR env var)",
    )
    parser.add_argument(
        "--resume-docx-input",
        default=DEFAULT_RESUME_DOCX_INPUT,
        help="Input DOCX resume template path. Default: Khaleel_Resume.docx (or RESUME_DOCX_INPUT env var)",
    )
    parser.add_argument(
        "--resume-docx-output",
        default=DEFAULT_TAILORED_RESUME_DOCX,
        help="Output tailored DOCX resume path. Default: khaleel_CV.docx (or RESUME_DOCX_OUTPUT env var)",
    )
    parser.add_argument(
        "--resume-doc-output",
        default=DEFAULT_TAILORED_RESUME_DOC,
        help="Output tailored DOC resume path. Default: khaleel_CV.doc (or RESUME_DOC_OUTPUT env var)",
    )
    parser.add_argument(
        "--k",
        type=int,
        default=8,
        help="Number of chunks to retrieve (default: 8)",
    )
    parser.add_argument(
        "--chunk-size",
        type=int,
        default=700,
        help="Chunk size used during indexing (default: 700)",
    )
    parser.add_argument(
        "--chunk-overlap",
        type=int,
        default=120,
        help="Chunk overlap used during indexing (default: 120)",
    )
    parser.add_argument(
        "--rebuild",
        action="store_true",
        help="Rebuild Chroma index even if existing index is present.",
    )
    parser.add_argument(
        "--env-file",
        dest="env_file",
        default=".env",
        help="Path to .env file (default: .env)",
    )
    parser.add_argument(
        "--model",
        default=None,
        help="Global model override (applies to ATS and cover letter if task-specific models are not set)",
    )
    parser.add_argument(
        "--ats-model",
        default=None,
        help="Model override for ATS scoring only",
    )
    parser.add_argument(
        "--cover-model",
        default=None,
        help="Model override for cover letter generation only",
    )
    parser.add_argument(
        "--provider",
        choices=["auto", "nvidia"],
        default="nvidia",
        help="Model provider to use (default: nvidia)",
    )
    parser.add_argument(
        "--generate-cover-letter",
        action="store_true",
        help="Deprecated no-op flag kept for backward compatibility.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    run_ats_and_cover_pipeline(args)


if __name__ == "__main__":
    main()
